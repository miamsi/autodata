"""
Budget Intelligence Agent - Executive Report Compiling Framework
"""

import pandas as pd
from typing import Dict, Any
from docx import Document
import io

def format_rp(value):
    try:
        is_neg = value < 0
        val = abs(value)
        fmt = f"{val:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")
        res = f"Rp {fmt}"
        return f"-{res}" if is_neg else res
    except:
        return "Rp 0"

def compile_executive_report(metrics: Dict[str, Any], forecast: Dict[str, Any], bottlenecks_df: pd.DataFrame, anomalies_df: pd.DataFrame, ai_recommendations: str) -> str:
    
    # Format the bottleneck dataframe so it looks good in Markdown (No scientific notation)
    b_display = bottlenecks_df.head(5).copy()
    for col in b_display.columns:
        if b_display[col].dtype in ['float64', 'int64']:
            if 'RATE' in col or 'SCORE' in col:
                b_display[col] = b_display[col].apply(lambda x: f"{x:.2f}")
            else:
                b_display[col] = b_display[col].apply(format_rp)
                
    top_bottlenecks = b_display.to_markdown(index=False)
    
    anomaly_summary_counts = {
        'Defisit': int(anomalies_df['ANOMALY_DEFICIT'].sum()),
        'Penyerapan Sangat Lambat': int(anomalies_df['ANOMALY_LOW_ABSORPTION'].sum()),
        'Lonjakan Ekstrem': int(anomalies_df['ANOMALY_SPIKE'].sum()),
        'Blokir Tinggi (>20%)': int(anomalies_df['ANOMALY_HIGH_BLOKIR'].sum())
    }
    
    report_md = f"""# RINGKASAN EKSEKUTIF INTELIJEN ANGGARAN
*Dihasilkan secara otomatis oleh Budget Intelligence Agent*

## 1. Tinjauan Makro Eksekutif
* **Total Alokasi Anggaran (Pagu):** {format_rp(metrics['total_budget'])}
* **Realisasi Tahun Berjalan (YTD):** {format_rp(metrics['total_realization'])} ({metrics['absorption_rate']:.2f}%)
* **Anggaran Diblokir Administratif:** {format_rp(metrics['total_blocked'])} ({metrics['blocked_rate']:.2f}%)

## 2. Proyeksi Kecepatan Penyerapan (Run-Rate) Akhir Tahun
* **Proyeksi Realisasi Akhir Tahun:** {format_rp(forecast['forecast_eoy_realization'])}
* **Proyeksi Persentase Penyerapan Akhir:** {forecast['forecast_absorption_rate']:.2f}%
* **Estimasi Risiko Gagal Serap (Under-Absorption):** {format_rp(forecast['under_absorption'])}
* **Indeks Risiko Penumpukan Desember:** {forecast['dec_concentration_pct']:.2f}% dari total belanja terkonsentrasi di akhir tahun.

## 3. Prioritas Hambatan Operasional (Top 5 Bottleneck)
Entitas di bawah ini menahan dana tunai/blokir terbesar dipadukan dengan persentase penyerapan yang lambat.

{top_bottlenecks}

## 4. Rangkuman Kerentanan & Anomali Data
* **Baris Anggaran Terindikasi Defisit:** {anomaly_summary_counts['Defisit']} baris
* **Baris Pagu Besar namun Sangat Lambat:** {anomaly_summary_counts['Penyerapan Sangat Lambat']} baris
* **Lonjakan Belanja Bulanan Abnormal:** {anomaly_summary_counts['Lonjakan Ekstrem']} baris
* **Dana Terkunci Blokir Berlebih:** {anomaly_summary_counts['Blokir Tinggi (>20%)']} baris

## 5. Rencana Aksi & Strategi Akselerasi (Saran AI)
{ai_recommendations}
---
*Akhir Laporan. Bersifat Rahasia (Confidential).*
"""
    return report_md

def generate_docx(markdown_text: str) -> bytes:
    """Mengubah teks Markdown menjadi file Microsoft Word (.docx)"""
    doc = Document()
    for line in markdown_text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), 0)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), 1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), 2)
        elif line.startswith('* '):
            doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
        elif line.strip() == '':
            continue
        else:
            doc.add_paragraph(line)
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
